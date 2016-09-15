#!/usr/bin/python
"""
dir2docx - convert files from directory and sub direcetories recursivly to a docx
"""

import os
import argparse
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

def log(txt):
    """ Print txt to stdout if verbose is TRUE """

    if ARGS.verbose:
        print(txt)

def generate_docx(path, outfile, lvl=1):
    """ Creates document and parses directory recursivly to generate docx file """

    document = Document()
    document = create_code_style(document)
 

    add_dir_to_dox(path, document, lvl)

    document.save(outfile)

def create_code_style(document):

    obj_styles = document.styles
    codestyle = obj_styles.add_style('Code_style', WD_STYLE_TYPE.PARAGRAPH)
    codestyle.base_style = obj_styles['Normal']
    codestyle.font.size = Pt(9)
    codestyle.font.name = 'Consolas'

    return document

def all_files_in_dir(dir_path):
    """ Returns non-hidden files in dir_path, if ARGS.all then all files are returned """

    dir_entries = sorted(os.listdir(dir_path))

    if not ARGS.all:
        dir_entries = [entry for entry in dir_entries if not entry.startswith('.')]

    return dir_entries

def add_dir_to_dox(dir_path, document, lvl):
    """ Adds all files in dir and enters directories recursively to do the same """

    log("Entering %s" % dir_path)

    dir_entries = all_files_in_dir(dir_path)

    for entry in dir_entries:
        path = os.path.join(dir_path, entry)

        if os.path.isdir(path):

            (_, title) = os.path.split(path)
            document.add_heading(title, lvl)

            add_dir_to_dox(path, document, lvl+1)
        else:
            add_file_to_docx(path, document, lvl)

def add_file_to_docx(path, document, lvl):
    """ Adds file to document if it is not in ignore list """

    extension = path.split(".")[-1]
    if extension in ARGS.ignore:
        log("Skipping %s" % path)
        return

    log("Adding %s" % path)

    title = os.path.basename(path)

    document.add_heading(title, lvl)

    if os.path.splitext(path)[-1] in ['.txt','.md','.markdown']:
        style = 'Normal'
    else:
        style = 'Code_style'

    log("> Applying {style} style for {ext}".format(ext=os.path.splitext(path), style=style))

    with open(path) as f:
        p = document.add_paragraph(f.read(), style=style)

def parse_args():
    """ Setup arguments"""

    parser = argparse.ArgumentParser()
    parser.add_argument("-a", "--all", action="store_true",
                        help="Parse all files and don't skip over hidden files & folders")
    parser.add_argument("-p", "--path", type=str,
                        default=".",
                        help="Top level directory path to scan")
    parser.add_argument("-o", "--outfile", type=str,
                        default="output.docx",
                        help="Name of output file")
    parser.add_argument("-i", "--ignore", nargs='+', type=str,
                        default=["exe"],
                        help="File extensions to ignore and not add to the docx")
    parser.add_argument("-v", "--verbose", action="store_true",
                        help="verbose")

    return parser.parse_args()


if __name__ == '__main__':

    ARGS = parse_args()
    generate_docx(ARGS.path, ARGS.outfile)

